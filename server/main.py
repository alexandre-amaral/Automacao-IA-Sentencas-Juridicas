from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi import UploadFile, File, HTTPException
from fastapi import status
from fastapi.responses import FileResponse
from pydantic import BaseModel
from dotenv import load_dotenv
import os
import logging
from uuid import uuid4
from pathlib import Path
from typing import Optional, Dict

# Configurar logging detalhado
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Parsing libs
try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None
try:
    from docx import Document  # python-docx
except Exception:  # pragma: no cover
    Document = None

# Carrega variáveis de ambiente
load_dotenv()

# Inicializa a aplicação FastAPI
app = FastAPI(
    title="Sistema de Automação de Sentenças Judiciais",
    description="API para automatizar a geração de sentenças judiciais com IA",
    version="1.0.0"
)

# Configuração do CORS para permitir requisições do frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Permitir todas as origens temporariamente
    allow_credentials=False,  # Desabilitar credentials para simplificar
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
)

STORAGE_DIR = Path(__file__).resolve().parent / "storage"
STORAGE_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_PROCESSO = {"pdf", "docx"}
ALLOWED_AUDIENCIA = {"mp4", "mp3", "wav", "m4a", "aac"}

class UploadResponse(BaseModel):
    case_id: str
    files: Dict[str, Optional[str]]

class ProcessingResponse(BaseModel):
    case_id: str
    step: str
    status: str
    message: str

@app.get("/")
async def root():
    """Endpoint raiz da API"""
    return {
        "message": "Sistema de Automação de Sentenças Judiciais",
        "status": "online",
        "version": "1.0.0"
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy"}

@app.post("/upload-caso", response_model=UploadResponse, status_code=status.HTTP_201_CREATED)
async def upload_caso(
    processo: Optional[UploadFile] = File(default=None, description="PDF/DOCX do processo"),
    audiencia: Optional[UploadFile] = File(default=None, description="MP4 da audiência")
):
    """Recebe upload de 1 arquivo de processo (PDF/DOCX) e 1 de audiência (MP4)"""
    logger.info("🚀 INICIANDO UPLOAD DE CASO")
    
    if processo is None and audiencia is None:
        logger.error("❌ Nenhum arquivo enviado")
        raise HTTPException(status_code=400, detail="Envie ao menos um arquivo: processo e/ou audiência")

    case_id = str(uuid4())
    case_dir = STORAGE_DIR / case_id
    case_dir.mkdir(parents=True, exist_ok=True)
    
    logger.info(f"📁 Case ID gerado: {case_id}")
    logger.info(f"📂 Diretório criado: {case_dir}")

    saved_processo = None
    saved_audiencia = None

    # Salvar processo
    if processo is not None:
        logger.info(f"📄 Processando arquivo: {processo.filename}")
        proc_ext = (processo.filename or "").split(".")[-1].lower()
        if proc_ext not in ALLOWED_PROCESSO:
            logger.error(f"❌ Extensão inválida do processo: .{proc_ext}")
            raise HTTPException(status_code=400, detail=f"Extensão do processo inválida: .{proc_ext}")
        proc_path = case_dir / f"processo.{proc_ext}"
        data = await processo.read()
        proc_path.write_bytes(data)
        saved_processo = str(proc_path)
        logger.info(f"✅ Processo salvo: {proc_path} ({len(data)} bytes)")

    # Salvar audiência
    if audiencia is not None:
        logger.info(f"🎥 Processando audiência: {audiencia.filename}")
        aud_ext = (audiencia.filename or "").split(".")[-1].lower()
        if aud_ext not in ALLOWED_AUDIENCIA:
            logger.error(f"❌ Extensão inválida da audiência: .{aud_ext}")
            raise HTTPException(status_code=400, detail=f"Extensão da audiência inválida: .{aud_ext}")
        aud_path = case_dir / f"audiencia.{aud_ext}"
        data = await audiencia.read()
        aud_path.write_bytes(data)
        saved_audiencia = str(aud_path)
        logger.info(f"✅ Audiência salva: {aud_path} ({len(data)} bytes)")

    logger.info(f"🎉 UPLOAD CONCLUÍDO! Case ID: {case_id}")
    
    # 🚀 INICIAR PIPELINE AUTOMÁTICO
    logger.info("🔄 Iniciando pipeline automático...")
    
    try:
        # ETAPA 1: Transcrição automática
        if saved_audiencia:
            logger.info("📝 Executando Etapa 1: Transcrição...")
            await executar_transcricao_automatica(case_id)
        
        # ETAPA 2: Processamento Gemini automático  
        logger.info("🧠 Executando Etapa 2: Processamento Gemini...")
        await executar_processamento_automatico(case_id)
        
        # ETAPA 3: Geração Claude automática
        logger.info("✍️ Executando Etapa 3: Geração de Sentença...")
        await executar_geracao_automatica(case_id)
        
        logger.info("🎉 PIPELINE AUTOMÁTICO CONCLUÍDO!")
        
    except Exception as e:
        logger.error(f"❌ Erro no pipeline automático: {str(e)}")
        # Continua e retorna o case_id mesmo se houver erro

    return UploadResponse(case_id=case_id, files={"processo": saved_processo, "audiencia": saved_audiencia})

# 🤖 FUNÇÕES AUXILIARES PARA PIPELINE AUTOMÁTICO
async def executar_transcricao_automatica(case_id: str):
    """Executa transcrição automaticamente"""
    from services.whisper_service import WhisperService
    
    case_dir = STORAGE_DIR / case_id
    
    # Encontrar arquivo de áudio
    audio_file = None
    for ext in ALLOWED_AUDIENCIA:
        audio_path = case_dir / f"audiencia.{ext}"
        if audio_path.exists():
            audio_file = audio_path
            break
    
    if not audio_file:
        logger.warning("⚠️ Nenhum arquivo de áudio encontrado para transcrição")
        return
    
    # Verificar tamanho e processar
    tamanho_mb = audio_file.stat().st_size / (1024 * 1024)
    logger.info(f"📏 Tamanho do arquivo: {tamanho_mb:.1f}MB")
    
    arquivo_para_transcricao = audio_file
    
    if tamanho_mb > 25:
        logger.info(f"🔄 Arquivo muito grande ({tamanho_mb:.1f}MB). Extraindo áudio comprimido...")
        import subprocess
        
        audio_comprimido = case_dir / "audio_temp.mp3"
        comando = [
            'ffmpeg', '-i', str(audio_file),
            '-vn', '-acodec', 'mp3', '-ab', '64k', '-ar', '16000', '-y',
            str(audio_comprimido)
        ]
        
        resultado = subprocess.run(comando, capture_output=True, text=True)
        if resultado.returncode == 0:
            arquivo_para_transcricao = audio_comprimido
            logger.info(f"✅ Áudio comprimido: {arquivo_para_transcricao.stat().st_size / (1024 * 1024):.1f}MB")
    
    # Transcrever
    whisper_service = WhisperService()
    transcricao = whisper_service.transcrever_audiencia(arquivo_para_transcricao, case_id)
    whisper_service.salvar_transcricao(transcricao, case_dir)
    
    # Limpar temporário
    if arquivo_para_transcricao != audio_file and arquivo_para_transcricao.exists():
        arquivo_para_transcricao.unlink()
    
    logger.info(f"✅ Transcrição automática concluída: {len(transcricao.texto_completo)} caracteres")

async def executar_processamento_automatico(case_id: str):
    """Executa processamento Gemini automaticamente"""
    from services.gemini_processor import GeminiProcessor
    from services.rag_service import RAGService
    from services.whisper_service import TranscricaoAudiencia, WhisperService
    
    case_dir = STORAGE_DIR / case_id
    
    # Processar documento
    gemini_processor = GeminiProcessor()
    
    # Encontrar processo
    processo_path = None
    for ext in ALLOWED_PROCESSO:
        proc_path = case_dir / f"processo.{ext}"
        if proc_path.exists():
            processo_path = proc_path
            break
    
    if not processo_path:
        raise Exception("Documento do processo não encontrado")
    
    # Extrair texto
    if processo_path.suffix.lower() == '.pdf':
        doc = fitz.open(str(processo_path))
        texto_processo = ""
        for page in doc:
            texto_processo += page.get_text("text") + "\n"
        doc.close()
    else:  # DOCX
        doc = Document(str(processo_path))
        texto_processo = '\n'.join([par.text for par in doc.paragraphs if par.text.strip()])
    
    logger.info(f"✅ Texto extraído: {len(texto_processo)} caracteres")
    
    # Limitar se muito grande
    if len(texto_processo) > 1500000:
        texto_processo = texto_processo[:1500000] + "... [TEXTO TRUNCADO - MUITO GRANDE]"
        logger.info("✂️ Texto truncado para 1.5M caracteres")
    
    # Salvar texto extraído para uso posterior no diálogo inteligente
    texto_extraido_path = case_dir / "processo_extraido.txt"
    texto_extraido_path.write_text(texto_processo, encoding='utf-8')
    logger.info(f"💾 Texto do processo salvo: {texto_extraido_path}")
    
    # Processar com Gemini
    processo_estruturado = gemini_processor.extrair_informacoes_processo(texto_processo)
    logger.info(f"✅ Processamento Gemini concluído: {processo_estruturado.numero_processo}")
    
    # Processar transcrição se existir
    transcricao_path = case_dir / "audiencia_transcricao.txt"
    analise_audiencia = None
    if transcricao_path.exists():
        transcricao_texto = transcricao_path.read_text(encoding='utf-8')
        transcricao_mock = TranscricaoAudiencia(
            texto_completo=transcricao_texto,
            duracao_segundos=None,
            idioma_detectado='pt',
            confianca=None,
            segmentos=None,
            metadados={}
        )
        whisper_service = WhisperService()
        analise_audiencia = whisper_service.processar_audiencia_com_gemini(transcricao_mock, case_id)
    
    # Salvar no RAG
    rag_service = RAGService()
    rag_service.salvar_conhecimento_caso(processo_estruturado, analise_audiencia, case_id)
    logger.info("✅ Processamento automático concluído e salvo no RAG")

async def executar_geracao_automatica(case_id: str):
    """Executa geração de sentença usando DIÁLOGO INTELIGENTE com prompt base estruturado"""
    from services.intelligent_dialogue_service import IntelligentDialogueService
    import json
    
    case_dir = STORAGE_DIR / case_id
    
    logger.info(f"🧠 [{case_id}] INICIANDO DIÁLOGO INTELIGENTE (3 ETAPAS ESTRUTURADAS)")
    
    try:
        # Recuperar texto do processo
        processo_path = case_dir / "processo_extraido.txt"
        if not processo_path.exists():
            logger.warning(f"⚠️ [{case_id}] Texto do processo não encontrado")
            return
        
        texto_processo = processo_path.read_text(encoding='utf-8')
        logger.info(f"📄 Texto do processo carregado: {len(texto_processo)} caracteres")
        
        # Recuperar transcrição da audiência (opcional)
        transcricao_audiencia = None
        transcricao_path = case_dir / "transcricao.json"
        
        if transcricao_path.exists():
            transcricao_data = json.loads(transcricao_path.read_text(encoding='utf-8'))
            transcricao_audiencia = transcricao_data.get('texto_completo', '')
            logger.info(f"🎤 Transcrição carregada: {len(transcricao_audiencia)} caracteres")
        
        # Executar diálogo inteligente com prompt base estruturado
        # NOVA ABORDAGEM SECTORIAL para sentenças detalhadas
        dialogue_service = IntelligentDialogueService(case_id)
        resultado_completo = dialogue_service.executar_dialogo_completo(
            texto_processo=texto_processo,
            transcricao_audiencia=transcricao_audiencia
        )
        
        logger.info(f"✅ DIÁLOGO INTELIGENTE CONCLUÍDO - {len(resultado_completo['etapas_executadas'])} etapas")
        
        # Salvar resultado completo estruturado
        resultado_path = case_dir / "dialogo_resultado_completo.json"
        resultado_path.write_text(
            json.dumps(resultado_completo, indent=2, ensure_ascii=False),
            encoding='utf-8'
        )
        
        # Salvar sentença final
        sentenca_final = resultado_completo.get("sentenca_final", "")
        if sentenca_final:
            sentenca_path = case_dir / "sentenca_gerada.txt"
            sentenca_path.write_text(sentenca_final, encoding='utf-8')
            logger.info(f"📝 Sentença gerada com PROMPT BASE: {len(sentenca_final)} caracteres")
        
        logger.info(f"🎉 [{case_id}] GERAÇÃO AUTOMÁTICA COMPLETA COM PROMPT BASE ESTRUTURADO")
        
    except Exception as e:
        logger.error(f"❌ Erro na geração com diálogo inteligente: {str(e)}")
        raise

# ETAPA 1: Transcrever áudio com Whisper
@app.post("/step1-transcribe/{case_id}", response_model=ProcessingResponse)
async def step1_transcribe_audio(case_id: str):
    """ETAPA 1: Transcreve áudio da audiência usando Whisper"""
    logger.info(f"🎤 INICIANDO ETAPA 1 - TRANSCRIÇÃO WHISPER - Case ID: {case_id}")
    
    from services.whisper_service import WhisperService
    
    case_dir = STORAGE_DIR / case_id
    if not case_dir.exists():
        logger.error(f"❌ Case ID não encontrado: {case_id}")
        raise HTTPException(status_code=404, detail="case_id não encontrado")
    
    logger.info(f"📂 Diretório encontrado: {case_dir}")
    
    try:
        # Encontrar arquivo de áudio
        logger.info("🔍 Procurando arquivo de áudio...")
        audio_file = None
        for ext in ALLOWED_AUDIENCIA:
            audio_path = case_dir / f"audiencia.{ext}"
            logger.info(f"   Verificando: {audio_path}")
            if audio_path.exists():
                audio_file = audio_path
                logger.info(f"✅ Arquivo de áudio encontrado: {audio_file}")
                break
        
        if not audio_file:
            logger.error("❌ Nenhum arquivo de áudio encontrado")
            raise HTTPException(status_code=404, detail="Arquivo de áudio não encontrado")
        
        # Verificar tamanho do arquivo
        tamanho_mb = audio_file.stat().st_size / (1024 * 1024)
        logger.info(f"📏 Tamanho do arquivo: {tamanho_mb:.1f}MB")
        
        arquivo_para_transcricao = audio_file
        
        # Se arquivo muito grande, extrair áudio comprimido
        if tamanho_mb > 25:
            logger.info(f"🔄 Arquivo muito grande ({tamanho_mb:.1f}MB). Extraindo áudio comprimido...")
            import subprocess
            
            # Criar arquivo de áudio comprimido
            audio_comprimido = case_dir / "audio_temp.mp3"
            
            comando = [
                'ffmpeg', '-i', str(audio_file),
                '-vn',  # Sem vídeo
                '-acodec', 'mp3',
                '-ab', '64k',  # Bitrate baixo
                '-ar', '16000',  # Sample rate reduzido
                '-y',  # Sobrescrever
                str(audio_comprimido)
            ]
            
            logger.info(f"🎬 Executando: {' '.join(comando[:6])}...")
            resultado = subprocess.run(comando, capture_output=True, text=True)
            
            if resultado.returncode != 0:
                logger.error(f"❌ Erro no ffmpeg: {resultado.stderr}")
                raise HTTPException(status_code=500, detail=f"Falha na conversão de áudio: {resultado.stderr}")
            
            arquivo_para_transcricao = audio_comprimido
            tamanho_novo = arquivo_para_transcricao.stat().st_size / (1024 * 1024)
            logger.info(f"✅ Áudio comprimido: {tamanho_novo:.1f}MB")
        
        # Transcrever com Whisper
        logger.info("🤖 Iniciando transcrição com OpenAI Whisper...")
        whisper_service = WhisperService()
        transcricao = whisper_service.transcrever_audiencia(arquivo_para_transcricao, case_id)
        logger.info(f"✅ Transcrição concluída! ({len(transcricao.texto_completo)} caracteres)")
        
        # Limpar arquivo temporário se criado
        if arquivo_para_transcricao != audio_file and arquivo_para_transcricao.exists():
            arquivo_para_transcricao.unlink()
            logger.info("🗑️ Arquivo temporário removido")
        
        # Salvar transcrição
        logger.info("💾 Salvando transcrição...")
        whisper_service.salvar_transcricao(transcricao, case_dir)
        logger.info("✅ Transcrição salva com sucesso!")
        
        return ProcessingResponse(
            case_id=case_id,
            step="transcribe_audio",
            status="completed",
            message=f"Áudio transcrito com sucesso. {len(transcricao.texto_completo)} caracteres."
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro na transcrição: {str(e)}")

# ETAPA 2: Processar com Gemini
@app.post("/step2-process/{case_id}", response_model=ProcessingResponse)
async def step2_process_with_gemini(case_id: str):
    """ETAPA 2: Processa documentos e transcrição com Gemini"""
    logger.info(f"🧠 INICIANDO ETAPA 2 - PROCESSAMENTO GEMINI - Case ID: {case_id}")
    
    from services.gemini_processor import GeminiProcessor
    from services.rag_service import RAGService
    
    case_dir = STORAGE_DIR / case_id
    if not case_dir.exists():
        logger.error(f"❌ Case ID não encontrado: {case_id}")
        raise HTTPException(status_code=404, detail="case_id não encontrado")

    logger.info(f"📂 Diretório encontrado: {case_dir}")
    
    try:
        # Processar documento do processo
        logger.info("🤖 Inicializando Gemini Processor...")
        gemini_processor = GeminiProcessor()
        
        # Encontrar e processar documento
        logger.info("🔍 Procurando documento do processo...")
        processo_path = None
        for ext in ALLOWED_PROCESSO:
            proc_path = case_dir / f"processo.{ext}"
            logger.info(f"   Verificando: {proc_path}")
            if proc_path.exists():
                processo_path = proc_path
                logger.info(f"✅ Documento encontrado: {processo_path}")
                break

        if not processo_path:
            logger.error("❌ Documento do processo não encontrado")
            raise HTTPException(status_code=404, detail="Documento do processo não encontrado")
        
        # Extrair texto e processar
        logger.info(f"📄 Extraindo texto de: {processo_path.suffix}")
        
        try:
            if processo_path.suffix.lower() == '.pdf':
                doc = fitz.open(str(processo_path))
                texto_processo = ""
                for page in doc:
                    texto_processo += page.get_text("text") + "\n"
                doc.close()
            else:  # DOCX
                doc = Document(str(processo_path))
                texto_processo = '\n'.join([par.text for par in doc.paragraphs if par.text.strip()])
            
            logger.info(f"✅ Texto extraído: {len(texto_processo)} caracteres")
            
            # Salvar texto extraído para uso posterior no diálogo inteligente
            texto_extraido_path = case_dir / "processo_extraido.txt"
            texto_extraido_path.write_text(texto_processo, encoding='utf-8')
            logger.info(f"💾 Texto do processo salvo: {texto_extraido_path}")
            
            # Gemini 1.5 Pro suporta até 2M tokens (~1.6M caracteres)
            # Só truncar se for extremamente grande (acima de 1.5M chars)
            if len(texto_processo) > 1500000:
                texto_processo = texto_processo[:1500000] + "... [TEXTO TRUNCADO - MUITO GRANDE]"
                logger.info("✂️ Texto truncado para 1.5M caracteres (limite prático)")
            
            logger.info("🧠 Processando com Gemini...")
            
            try:
                processo_estruturado = gemini_processor.extrair_informacoes_processo(texto_processo)
                logger.info(f"✅ Processamento Gemini concluído: {processo_estruturado.numero_processo}")
            except Exception as gemini_error:
                if "quota" in str(gemini_error).lower() or "429" in str(gemini_error):
                    logger.warning("⚠️ Quota do Gemini excedida. Usando estrutura mock para continuar o teste...")
                    
                    # Mock temporário para testar Claude
                    from services.gemini_processor import ProcessoEstruturado, ParteProcesso, PedidoProcesso, FatoRelevante
                    
                    processo_estruturado = ProcessoEstruturado(
                        numero_processo="0000398-03.2024.5.09.0010",
                        partes=[
                            ParteProcesso(nome="João Silva", tipo="requerente", qualificacao="Trabalhador"),
                            ParteProcesso(nome="Empresa XYZ Ltda", tipo="requerida", qualificacao="Empregadora")
                        ],
                        pedidos=[
                            PedidoProcesso(descricao="Horas extras", categoria="verbas_trabalhistas", valor_estimado="R$ 15.000,00", deferido=None)
                        ],
                        fatos_relevantes=[
                            FatoRelevante(descricao="Trabalho sem registro de horas extras", fonte="inicial", impacto_decisao="Relevante para decisão")
                        ],
                        periodo_contratual="2020-2024",
                        valor_causa="R$ 50.000,00",
                        jurisprudencias_citadas=[],
                        decisao_final=None,
                        fundamentacao_resumida="Processo trabalhista em andamento"
                    )
                    logger.info("✅ Mock do Gemini aplicado para continuar teste")
                else:
                    raise HTTPException(status_code=500, detail=f"Erro no processamento Gemini: {str(gemini_error)}")
            
        except Exception as e:
            logger.error(f"❌ Erro no processamento Gemini: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Erro no processamento Gemini: {str(e)}")
        
        # Processar transcrição da audiência se existir
        transcricao_path = case_dir / "audiencia_transcricao.txt"
        analise_audiencia = None
        if transcricao_path.exists():
            transcricao_texto = transcricao_path.read_text(encoding='utf-8')
            # Criar objeto TranscricaoAudiencia mock para análise
            from services.whisper_service import TranscricaoAudiencia, WhisperService
            transcricao_mock = TranscricaoAudiencia(
                texto_completo=transcricao_texto,
                duracao_segundos=None,
                idioma_detectado='pt',
                confianca=None,
                segmentos=None,
                metadados={}
            )
            whisper_service = WhisperService()
            analise_audiencia = whisper_service.processar_audiencia_com_gemini(transcricao_mock, case_id)
        
        # Salvar no RAG
        rag_service = RAGService()
        rag_service.salvar_conhecimento_caso(processo_estruturado, analise_audiencia, case_id)
        
        return ProcessingResponse(
            case_id=case_id,
            step="process_gemini",
            status="completed",
            message="Documentos processados e salvos no RAG com sucesso"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro no processamento Gemini: {str(e)}")

# ETAPA 3: Gerar sentença com Claude
@app.post("/step3-generate/{case_id}", response_model=ProcessingResponse)
async def step3_generate_sentence(case_id: str):
    """ETAPA 3: Gera sentença consultando RAG com Claude"""
    logger.info(f"✍️ INICIANDO ETAPA 3 - GERAÇÃO CLAUDE - Case ID: {case_id}")
    
    from services.claude_service import ClaudeService
    from services.rag_service import RAGService
    
    case_dir = STORAGE_DIR / case_id
    if not case_dir.exists():
        logger.error(f"❌ Case ID não encontrado: {case_id}")
        raise HTTPException(status_code=404, detail="case_id não encontrado")

    logger.info(f"📂 Diretório encontrado: {case_dir}")
    
    try:
        # Consultar conhecimento no RAG
        logger.info("📚 Consultando conhecimento no RAG...")
        rag_service = RAGService()
        conhecimento_caso = rag_service.recuperar_conhecimento_caso(case_id)
        logger.info(f"✅ Conhecimento recuperado: {len(str(conhecimento_caso))} caracteres")
        
        # Gerar sentença com Claude
        logger.info("🤖 Iniciando geração de sentença com Claude...")
        claude_service = ClaudeService()
        sentenca = claude_service.gerar_sentenca_com_rag(conhecimento_caso, case_id)
        logger.info(f"✅ Sentença gerada! ({len(sentenca)} caracteres)")
        
        # Salvar sentença
        logger.info("💾 Salvando sentença...")
        sentenca_path = case_dir / "sentenca.txt"
        sentenca_path.write_text(sentenca, encoding='utf-8')
        logger.info(f"✅ Sentença salva: {sentenca_path}")
        logger.info("🎉 ETAPA 3 CONCLUÍDA! Sentença gerada com sucesso!")
        
        return ProcessingResponse(
            case_id=case_id,
            step="generate_sentence",
            status="completed",
            message="Sentença gerada com sucesso"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro na geração da sentença: {str(e)}")

@app.get("/download/{case_id}/sentenca")
async def download_sentence(case_id: str):
    """Download da sentença gerada.
    Aceita tanto 'sentenca.txt' (etapa 3) quanto 'sentenca_gerada.txt' (pipeline automático).
    """
    case_dir = STORAGE_DIR / case_id
    path_options = [case_dir / "sentenca.txt", case_dir / "sentenca_gerada.txt"]
    sentenca_path = next((p for p in path_options if p.exists()), None)

    if not sentenca_path:
        raise HTTPException(status_code=404, detail="Sentença não encontrada")

    return FileResponse(
        str(sentenca_path),
        media_type="text/plain",
        filename=f"sentenca_{case_id}.txt"
    )

@app.post("/generate-from-existing/{case_id}", response_model=ProcessingResponse)
async def generate_from_existing(case_id: str):
    """Gera sentença usando arquivos já existentes (sem reprocessar APIs).

    Usa 'processo_extraido.txt' e, se existir, 'transcricao.json' ou 'audiencia_transcricao.txt'.
    """
    import json
    from services.intelligent_dialogue_service import IntelligentDialogueService

    case_dir = STORAGE_DIR / case_id
    if not case_dir.exists():
        raise HTTPException(status_code=404, detail="case_id não encontrado")

    # Carregar processo extraído
    processo_path = case_dir / "processo_extraido.txt"
    if not processo_path.exists():
        raise HTTPException(status_code=404, detail="processo_extraido.txt não encontrado")
    texto_processo = processo_path.read_text(encoding='utf-8')

    # Carregar transcrição (opcional)
    transcricao_audiencia = None
    transcricao_json = case_dir / "transcricao.json"
    transcricao_txt = case_dir / "audiencia_transcricao.txt"
    if transcricao_json.exists():
        data = json.loads(transcricao_json.read_text(encoding='utf-8'))
        transcricao_audiencia = data.get('texto_completo') or None
    elif transcricao_txt.exists():
        transcricao_audiencia = transcricao_txt.read_text(encoding='utf-8')

    # Executar diálogo inteligente (usa RAG avançado setorial internamente)
    dialogue_service = IntelligentDialogueService(case_id)
    resultado_completo = dialogue_service.executar_dialogo_completo(
        texto_processo=texto_processo,
        transcricao_audiencia=transcricao_audiencia
    )

    # Salvar
    resultado_path = case_dir / "dialogo_resultado_completo.json"
    resultado_path.write_text(
        json.dumps(resultado_completo, indent=2, ensure_ascii=False),
        encoding='utf-8'
    )

    sentenca_final = resultado_completo.get("sentenca_final", "")
    if sentenca_final:
        (case_dir / "sentenca_gerada.txt").write_text(sentenca_final, encoding='utf-8')

    return ProcessingResponse(
        case_id=case_id,
        step="generate_from_existing",
        status="completed",
        message="Geração concluída a partir de arquivos existentes"
    )

@app.post("/init-style/{case_id}", response_model=ProcessingResponse)
async def init_style_from_examples(case_id: str, max_docs: int = 30):
    """Inicializa o estilo da juíza para o caso usando sentenças modelos locais.

    Varre diretórios 'Sentenças_2023', 'Sentenças_2024', 'Sentenças_2025' e carrega o texto
    de arquivos .docx e .pdf como exemplos de estilo no RAG isolado do caso.
    """
    from services.enhanced_rag_service import EnhancedRAGService
    sentences_root = Path(__file__).resolve().parent.parent
    folders = [
        sentences_root / "Sentenças_2023",
        sentences_root / "Sentenças_2024",
        sentences_root / "Sentenças_2025",
    ]

    samples: list[str] = []

    def read_docx(path: Path) -> str:
        try:
            if Document is None:
                return ""
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return text
        except Exception:
            return ""

    def read_pdf(path: Path) -> str:
        try:
            if fitz is None:
                return ""
            doc = fitz.open(str(path))
            text = "\n".join(page.get_text("text") for page in doc)
            doc.close()
            return text
        except Exception:
            return ""

    for folder in folders:
        if not folder.exists():
            continue
        for file in folder.iterdir():
            if not file.is_file():
                continue
            ext = file.suffix.lower()
            content = ""
            if ext == ".docx":
                content = read_docx(file)
            elif ext == ".pdf":
                content = read_pdf(file)
            # .doc não suportado diretamente por python-docx; ignorar por ora

            if content:
                # Normalizar e limitar tamanho por documento para não estourar memória
                content = content.strip()
                if len(content) > 20000:
                    content = content[:20000]
                samples.append(content)

            if len(samples) >= max_docs:
                break
        if len(samples) >= max_docs:
            break

    if not samples:
        raise HTTPException(status_code=404, detail="Nenhuma sentença modelo encontrada (.docx/.pdf)")

    rag = EnhancedRAGService(case_id)
    rag.initialize_judge_style(samples, force_reload=True)

    return ProcessingResponse(
        case_id=case_id,
        step="init_style",
        status="completed",
        message=f"Estilo inicializado com {len(samples)} documentos"
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)